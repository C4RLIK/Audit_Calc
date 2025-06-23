import pandas as pd
import numpy as np
import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt

# Функция для расчёта уровня существенности
def calculate_materiality(data, deviation_threshold, rounding_limit):
    try:
        # Получаем значения показателей
        values = data['Значение'].dropna().astype(float).values
        if len(values) == 0:
            return None, "Нет данных для расчёта"
        
        # Рассчитываем среднее и отклонения
        mean = np.mean(values)
        deviations = [(x, abs(x - mean) / mean * 100) for x in values]
        
        # Фильтруем показатели по допустимому отклонению
        filtered = [x for x, dev in deviations if dev <= deviation_threshold]
        excluded = [x for x, dev in deviations if dev > deviation_threshold]
        
        if not filtered:
            return None, "Все показатели исключены как нерепрезентативные"
        
        # Рассчитываем новое среднее
        new_mean = np.mean(filtered)
        
        # Округляем результат
        rounded = round(new_mean / 100) * 100
        if abs(rounded - new_mean) > rounding_limit:
            rounded = new_mean
        
        # Формируем детали расчёта
        details = {
            "Исходные данные": data,
            "Среднее арифметическое": mean,
            "Отклонения": deviations,
            "Исключённые": excluded,
            "Оставшиеся": filtered,
            "Новое среднее": new_mean,
            "Округлённое": rounded,
            "Показатели": data['Показатель'].values
        }
        
        return rounded, details
    
    except Exception as e:
        return None, f"Ошибка расчёта: {str(e)}"

# Функция для создания Word-отчёта
def create_word_report(details, deviation_threshold):
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Заголовок
    title = doc.add_heading('Расчёт уровня существенности', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Исходные данные
    doc.add_heading('1. Исходные данные:', level=2)
    for idx, (name, value) in enumerate(zip(details["Показатели"], details["Исходные данные"]['Значение']), 1):
        doc.add_paragraph(f"{idx}. {name}: {value:,.0f} руб.", style='ListNumber')
    
    # 2. Среднее арифметическое
    doc.add_heading('2. Расчёт среднего арифметического:', level=2)
    values_str = " + ".join([f"{x:,.0f}" for x in details["Исходные данные"]['Значение']])
    doc.add_paragraph(f"({values_str}) / {len(details['Исходные данные'])} = {details['Среднее арифметическое']:,.0f} руб.")
    
    # 3. Отклонения показателей
    doc.add_heading('3. Определение отклонений показателей от среднего:', level=2)
    for x, dev in details["Отклонения"]:
        doc.add_paragraph(f"• {(x - details['Среднее арифметическое'])/details['Среднее арифметическое']*100:+.2f}%", style='ListBullet')
    
    # 4. Исключение нерепрезентативных показателей
    doc.add_heading(f'4. Исключение показателей с отклонением > {deviation_threshold}%:', level=2)
    if details["Исключённые"]:
        for x in details["Исключённые"]:
            doc.add_paragraph(f"• {x:,.0f} руб.", style='ListBullet')
    else:
        doc.add_paragraph("Нет исключённых показателей")
    
    # 5. Новое среднее
    doc.add_heading('5. Расчёт нового среднего арифметического:', level=2)
    doc.add_paragraph(f"({' + '.join([f'{x:,.0f}' for x in details['Оставшиеся']])}) / {len(details['Оставшиеся'])} = {details['Новое среднее']:,.2f} руб.")
    
    # 6. Округление
    doc.add_heading('6. Округление результата:', level=2)
    doc.add_paragraph(f"Округлённое значение: {details['Округлённое']:,.0f} руб.")
    
    # 7. Итог
    doc.add_heading('7. Итоговый уровень существенности:', level=2)
    p = doc.add_paragraph()
    p.add_run(f"{details['Округлённое']:,.0f} рублей").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

# Функция для отображения отчёта в Streamlit
def display_report(details, deviation_threshold, result):
    st.markdown("## Полный отчёт о расчёте уровня существенности")
    
    with st.expander("1. Исходные данные", expanded=True):
        for idx, (name, value) in enumerate(zip(details["Показатели"], details["Исходные данные"]['Значение']), 1):
            st.markdown(f"{idx}. **{name}**: {value:,.0f} руб.")
    
    with st.expander("2. Среднее арифметическое", expanded=True):
        values_str = " + ".join([f"{x:,.0f}" for x in details["Исходные данные"]['Значение']])
        st.markdown(f"({values_str}) / {len(details['Исходные данные'])} = **{details['Среднее арифметическое']:,.0f} руб.**")
    
    with st.expander("3. Отклонения показателей", expanded=True):
        for x, dev in details["Отклонения"]:
            st.markdown(f"- {x:,.0f} руб.: отклонение **{dev:+.2f}%**")
    
    with st.expander(f"4. Исключение показателей (отклонение > {deviation_threshold}%)", expanded=True):
        if details["Исключённые"]:
            for x in details["Исключённые"]:
                st.markdown(f"- {x:,.0f} руб.")
        else:
            st.markdown("Нет исключённых показателей")
    
    with st.expander("5. Новое среднее арифметическое", expanded=True):
        st.markdown(f"({' + '.join([f'{x:,.0f}' for x in details['Оставшиеся']])}) / {len(details['Оставшиеся'])} = **{details['Новое среднее']:,.2f} руб.**")
    
    with st.expander("6. Округление результата", expanded=True):
        st.markdown(f"Округлённое значение: **{details['Округлённое']:,.0f} руб.**")
    
    st.success(f"## Итоговый уровень существенности: {details['Округлённое']:,.0f} рублей")
    
    # Визуализация
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.bar(["Все показатели", "После исключения"], 
           [len(details["Исходные данные"]), len(details["Оставшиеся"])],
           color=["lightcoral", "lightgreen"])
    ax.set_title("Количество показателей до и после исключения")
    st.pyplot(fig)

# Основная функция
def main():
    st.set_page_config(
        page_title="Аудит: Уровень существенности",
        page_icon="🧮",
        layout="wide"
    )
    
    st.title("🧮 Калькулятор уровня существенности")
    st.markdown("""
    ### Инструкция:
    1. Выберите способ ввода данных (файл Excel или ручной ввод)
    2. Укажите параметры расчёта
    3. Нажмите кнопку "Рассчитать"
    4. Просмотрите отчёт и скачайте результаты
    """)
    
    # Выбор способа ввода данных
    input_method = st.radio(
        "**Выберите способ ввода данных:**",
        ["📊 Загрузить файл Excel", "✏️ Ввести вручную"],
        horizontal=True
    )
    
    data = None
    
    # Вариант 1: Загрузка файла Excel
    if input_method == "📊 Загрузить файл Excel":
        uploaded_file = st.file_uploader(
            "Загрузите файл Excel с показателями (столбцы: 'Показатель', 'Значение')", 
            type=["xlsx", "xls"]
        )
        
        if uploaded_file:
            try:
                data = pd.read_excel(uploaded_file, engine='openpyxl')
                if not all(col in data.columns for col in ['Показатель', 'Значение']):
                    st.error("Файл должен содержать столбцы 'Показатель' и 'Значение'")
                    data = None
                else:
                    st.success("Файл успешно загружен!")
                    st.dataframe(data)
            except Exception as e:
                st.error(f"Ошибка загрузки файла: {e}")
    
    # Вариант 2: Ручной ввод
    else:
        st.markdown("### Введите показатели вручную")
        num_indicators = st.number_input(
            "Количество показателей", 
            min_value=1, max_value=20, value=5, step=1
        )
        
        indicators = []
        for i in range(num_indicators):
            cols = st.columns(2)
            with cols[0]:
                name = st.text_input(
                    f"Название показателя {i+1}", 
                    value=f"Показатель {i+1}",
                    key=f"name_{i}"
                )
            with cols[1]:
                value = st.number_input(
                    f"Значение {i+1}", 
                    min_value=0, 
                    value=(i+1)*100000, 
                    step=1000,
                    key=f"value_{i}"
                )
            indicators.append({"Показатель": name, "Значение": value})
        
        if indicators:
            data = pd.DataFrame(indicators)
            st.dataframe(data)
    
    # Настройки расчёта
    st.sidebar.header("⚙️ Параметры расчёта")
    deviation = st.sidebar.slider(
        "Допустимое отклонение от среднего (%)", 
        min_value=0, max_value=100, value=50, step=1
    )
    rounding_limit = st.sidebar.number_input(
        "Максимальное отклонение при округлении", 
        min_value=0, value=50, step=10
    )
    
    # Кнопка расчёта
    if st.button("Рассчитать", type="primary") and data is not None:
        with st.spinner("Выполняется расчёт..."):
            result, details = calculate_materiality(data, deviation, rounding_limit)
        
        if result is None:
            st.error(details)
        else:
            # Отображение отчёта
            display_report(details, deviation, result)
            
            # Генерация и скачивание Word-отчёта
            doc = create_word_report(details, deviation)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            st.download_button(
                label="📥 Скачать отчёт в Word",
                data=buffer,
                file_name="Уровень_существенности.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                help="Нажмите для скачивания полного отчёта в формате Word"
            )

if __name__ == "__main__":
    main()
